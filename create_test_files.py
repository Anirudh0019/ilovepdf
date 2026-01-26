"""Generate test files for PDF tools testing"""
from pathlib import Path
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.colors import red, blue, green, black, purple
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
from reportlab.pdfgen import canvas
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

TESTS_DIR = Path("tests")
TESTS_DIR.mkdir(exist_ok=True)


def create_sample_pdf_1():
    """Create a simple 2-page PDF document"""
    output_path = TESTS_DIR / "document_1.pdf"
    doc = SimpleDocTemplate(str(output_path), pagesize=letter)
    styles = getSampleStyleSheet()
    story = []

    # Page 1
    story.append(Paragraph("Sample Document 1", styles["Title"]))
    story.append(Spacer(1, 0.5 * inch))
    story.append(Paragraph(
        "This is the first test PDF document. It contains sample text that you can use "
        "to test the PDF tools. This document has 2 pages.",
        styles["Normal"]
    ))
    story.append(Spacer(1, 0.25 * inch))
    story.append(Paragraph(
        "Lorem ipsum dolor sit amet, consectetur adipiscing elit. Sed do eiusmod tempor "
        "incididunt ut labore et dolore magna aliqua. Ut enim ad minim veniam, quis nostrud "
        "exercitation ullamco laboris nisi ut aliquip ex ea commodo consequat.",
        styles["Normal"]
    ))
    story.append(Spacer(1, 0.25 * inch))

    # Add a simple table
    data = [
        ["Feature", "Status"],
        ["Merge PDF", "Ready"],
        ["Split PDF", "Ready"],
        ["Compress", "Ready"],
        ["Watermark", "Ready"],
    ]
    table = Table(data, colWidths=[2 * inch, 1.5 * inch])
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), blue),
        ("TEXTCOLOR", (0, 0), (-1, 0), "white"),
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, 0), 12),
        ("BOTTOMPADDING", (0, 0), (-1, 0), 12),
        ("BACKGROUND", (0, 1), (-1, -1), "lightgrey"),
        ("GRID", (0, 0), (-1, -1), 1, black),
    ]))
    story.append(table)

    # Page 2
    story.append(PageBreak())
    story.append(Paragraph("Page 2 - Continued", styles["Heading1"]))
    story.append(Spacer(1, 0.5 * inch))
    story.append(Paragraph(
        "This is the second page of the document. When you split this PDF, "
        "you should get two separate files.",
        styles["Normal"]
    ))
    story.append(Spacer(1, 0.25 * inch))
    for i in range(1, 6):
        story.append(Paragraph(f"• Bullet point number {i}", styles["Normal"]))

    doc.build(story)
    print(f"Created: {output_path}")


def create_sample_pdf_2():
    """Create another PDF for merging tests"""
    output_path = TESTS_DIR / "document_2.pdf"
    doc = SimpleDocTemplate(str(output_path), pagesize=letter)
    styles = getSampleStyleSheet()
    story = []

    story.append(Paragraph("Sample Document 2", styles["Title"]))
    story.append(Spacer(1, 0.5 * inch))
    story.append(Paragraph(
        "This is the SECOND test PDF. Use this along with document_1.pdf to test "
        "the merge functionality.",
        styles["Normal"]
    ))
    story.append(Spacer(1, 0.25 * inch))
    story.append(Paragraph(
        "After merging, you should see both documents combined into a single PDF file.",
        styles["Normal"]
    ))
    story.append(Spacer(1, 0.5 * inch))

    # Add some colored text
    red_style = ParagraphStyle("RedText", parent=styles["Normal"], textColor=red)
    blue_style = ParagraphStyle("BlueText", parent=styles["Normal"], textColor=blue)
    green_style = ParagraphStyle("GreenText", parent=styles["Normal"], textColor=green)

    story.append(Paragraph("This text is RED", red_style))
    story.append(Paragraph("This text is BLUE", blue_style))
    story.append(Paragraph("This text is GREEN", green_style))

    doc.build(story)
    print(f"Created: {output_path}")


def create_sample_pdf_3():
    """Create a 5-page PDF for split testing"""
    output_path = TESTS_DIR / "multi_page.pdf"
    doc = SimpleDocTemplate(str(output_path), pagesize=A4)
    styles = getSampleStyleSheet()
    story = []

    for page_num in range(1, 6):
        story.append(Paragraph(f"Page {page_num} of 5", styles["Title"]))
        story.append(Spacer(1, 0.5 * inch))
        story.append(Paragraph(
            f"This is page number {page_num}. Use this document to test the split "
            f"functionality. Try splitting pages '1,3,5' or '2-4' to extract specific pages.",
            styles["Normal"]
        ))
        story.append(Spacer(1, 0.25 * inch))
        story.append(Paragraph(
            f"Content unique to page {page_num}: {'*' * (page_num * 10)}",
            styles["Normal"]
        ))
        if page_num < 5:
            story.append(PageBreak())

    doc.build(story)
    print(f"Created: {output_path}")


def create_sample_images():
    """Create test images for resize/compress testing"""
    # Create a colorful test image
    img1 = Image.new("RGB", (800, 600), color=(30, 30, 60))
    draw = ImageDraw.Draw(img1)

    # Draw some shapes
    draw.rectangle([50, 50, 200, 200], fill=(255, 100, 100))
    draw.ellipse([250, 50, 450, 200], fill=(100, 255, 100))
    draw.polygon([(550, 200), (650, 50), (750, 200)], fill=(100, 100, 255))

    # Add text
    draw.text((300, 300), "Test Image 1", fill=(255, 255, 255))
    draw.text((250, 350), "800 x 600 pixels", fill=(200, 200, 200))
    draw.text((200, 450), "Use for resize/compress testing", fill=(150, 150, 150))

    img1.save(TESTS_DIR / "test_image_large.png")
    img1.save(TESTS_DIR / "test_image_large.jpg", quality=95)
    print(f"Created: tests/test_image_large.png")
    print(f"Created: tests/test_image_large.jpg")

    # Create a smaller image
    img2 = Image.new("RGB", (400, 300), color=(60, 30, 60))
    draw2 = ImageDraw.Draw(img2)
    draw2.rectangle([20, 20, 380, 280], outline=(255, 255, 255), width=3)
    draw2.text((120, 130), "Small Image", fill=(255, 255, 255))
    draw2.text((130, 160), "400 x 300", fill=(200, 200, 200))
    img2.save(TESTS_DIR / "test_image_small.png")
    print(f"Created: tests/test_image_small.png")

    # Create an image with transparency (PNG)
    img3 = Image.new("RGBA", (500, 500), color=(0, 0, 0, 0))
    draw3 = ImageDraw.Draw(img3)
    draw3.ellipse([50, 50, 450, 450], fill=(255, 100, 50, 200))
    draw3.text((180, 230), "Transparent BG", fill=(255, 255, 255, 255))
    img3.save(TESTS_DIR / "test_transparent.png")
    print(f"Created: tests/test_transparent.png")


def create_sample_word_doc():
    """Create a Word document for conversion testing"""
    doc = Document()

    # Title
    title = doc.add_heading("Sample Word Document", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Introduction
    doc.add_paragraph(
        "This is a sample Word document created for testing the Word to PDF conversion feature. "
        "The converter should preserve basic text formatting and structure."
    )

    # Heading and content
    doc.add_heading("Section 1: Features", level=1)
    doc.add_paragraph(
        "The local PDF toolkit includes several useful features for handling documents privately:"
    )

    # Bullet points
    for feature in ["Merge multiple PDFs", "Split PDF by pages", "Compress PDF files",
                    "Add watermarks", "Convert documents"]:
        doc.add_paragraph(feature, style="List Bullet")

    doc.add_heading("Section 2: Privacy", level=1)
    doc.add_paragraph(
        "All files are processed locally on your machine. No data is sent to external servers. "
        "Files are immediately deleted after processing is complete."
    )

    # Another heading
    doc.add_heading("Section 3: Usage Notes", level=2)
    para = doc.add_paragraph()
    para.add_run("Important: ").bold = True
    para.add_run("This converter handles basic text content well. Complex formatting like tables, "
                 "images, and advanced styling may not be perfectly preserved.")

    doc.save(TESTS_DIR / "sample_document.docx")
    print(f"Created: tests/sample_document.docx")


if __name__ == "__main__":
    print("Creating test files...\n")

    create_sample_pdf_1()
    create_sample_pdf_2()
    create_sample_pdf_3()
    create_sample_images()
    create_sample_word_doc()

    print("\n✓ All test files created in tests/ directory!")
    print("\nTest scenarios:")
    print("  • Merge: document_1.pdf + document_2.pdf")
    print("  • Split: multi_page.pdf (try pages '1,3,5' or '2-4')")
    print("  • Compress: any PDF file")
    print("  • Resize: test_image_large.png (try 200x200)")
    print("  • Compress Image: test_image_large.jpg")
    print("  • Watermark: any PDF file")
    print("  • Word to PDF: sample_document.docx")
    print("  • PDF to Images: multi_page.pdf")
